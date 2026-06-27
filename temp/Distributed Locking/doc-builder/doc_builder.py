#!/usr/bin/env python3
"""Rendering engine: clone the teacher's template, strip its body, then emit new
content using the template's own named styles so the result is visually identical
in formatting (bilingual paragraphs, boxed quotes, captions, review questions).
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = os.path.dirname(__file__)
TEMPLATE = os.path.join(HERE, "..", "Fundamentals of Project Management Viet Anh.docx")
IMAGES = os.path.join(HERE, "images")

HL = {"yellow": WD_COLOR_INDEX.YELLOW, "green": WD_COLOR_INDEX.BRIGHT_GREEN,
      "cyan": WD_COLOR_INDEX.TURQUOISE}


class DocBuilder:
    def __init__(self):
        self.doc = Document(TEMPLATE)
        self._strip_body()
        self.fig_no = {}  # chapter -> running figure index

    def _strip_body(self):
        body = self.doc.element.body
        for child in list(body):
            if child.tag == qn("w:sectPr"):
                continue
            body.remove(child)

    # -- primitives ---------------------------------------------------------
    def _p(self, text, style, highlight=None):
        p = self.doc.add_paragraph(style=style)
        run = p.add_run(text)
        if highlight:
            run.font.highlight_color = HL[highlight]
        return p

    # -- front matter: cover page + table of contents -----------------------
    def _centered(self, text, size=None, bold=True, space_after=6):
        p = self.doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        return p

    def _blank(self, n=1):
        for _ in range(n):
            self.doc.add_paragraph(style="Normal")

    def _page_break(self):
        p = self.doc.add_paragraph()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        p.add_run()._r.append(br)

    def cover_page(self, title_en, title_vi, subject="CHUYÊN ĐỀ",
                   members=None, instructor="", year="2026"):
        self._centered("ĐẠI HỌC QUỐC GIA THÀNH PHỐ HỒ CHÍ MINH", size=14)
        self._centered("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN", size=14)
        self._blank(5)
        self._centered(subject, size=30)
        self._centered(title_en, size=30)
        self._centered(f"({title_vi})", size=18, bold=False)
        self._blank(4)
        self._centered("Giảng viên hướng dẫn: " + (instructor or "……………………"),
                       size=14, bold=False, space_after=4)
        self._centered("Nhóm thực hiện:", size=14, bold=False, space_after=4)
        for m in (members or ["…………………………", "…………………………", "…………………………"]):
            self._centered(m, size=13, bold=False, space_after=2)
        self._blank(3)
        self._centered(f"TP. Hồ Chí Minh – {year}", size=14, bold=False)
        self._page_break()

    def _toc_field(self):
        """Insert a real Word TOC field (levels 1-3). Word fills it in with the
        correct page numbers when the document is opened / fields are updated."""
        p = self.doc.add_paragraph()
        r = p.add_run()._r
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        begin.set(qn("w:dirty"), "true")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = ("Nhấn Ctrl + Click vào mục lục hoặc chọn rồi bấm F9 "
                            "để cập nhật số trang.")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for el in (begin, instr, sep, placeholder, end):
            r.append(el)

    def table_of_contents(self):
        self._centered("MỤC LỤC", size=16)
        self._toc_field()
        self._page_break()

    def _enable_update_fields(self):
        """Make Word refresh field values (the TOC) when the file is opened."""
        settings = self.doc.settings.element
        if settings.find(qn("w:updateFields")) is None:
            el = OxmlElement("w:updateFields")
            el.set(qn("w:val"), "true")
            settings.append(el)

    # -- headings -----------------------------------------------------------
    def h1(self, vi, en):
        self._p(f"{vi} – {en}", "Heading 1")

    def h2(self, vi, en):
        self._p(f"{vi} – {en}", "Heading 2")

    def h3(self, vi, en):
        self._p(f"{vi} – {en}", "Heading 3")

    # -- body paragraph pair ------------------------------------------------
    def p(self, en, vi):
        self._p(en, "Normal")
        self._p(vi, "NormalViet")

    # -- boxed important quote ---------------------------------------------
    def imp(self, en, vi, highlight="yellow"):
        self._p(en, "13. Importance")
        self._p(vi, "13. ImportanceViet", highlight=highlight)

    # -- figure: English image + caption, then Vietnamese image + caption ---
    def fig(self, chapter, base, cap_vi, cap_en):
        n = self.fig_no.get(chapter, 0) + 1
        self.fig_no[chapter] = n
        tag = f"{chapter}-{n}"
        self._image(f"{base}_en.png")
        self._p(f"Figure {tag}. {cap_en}", "08. Hình Ảnh")
        self._image(f"{base}_vi.png")
        self._p(f"Hình {tag}. {cap_vi}", "08. Hình Ảnh")

    def _image(self, filename, width=5.9):
        path = os.path.join(IMAGES, filename)
        self.doc.add_picture(path, width=Inches(width))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -- key points ---------------------------------------------------------
    def key_points(self, items):
        self._p("Những điểm chính cần nhớ trong chương – Key Points to Remember",
                "Heading 2")
        for vi in items:
            self._p(vi, "02. GachCongDauDongViet")

    # -- per-section review questions block ---------------------------------
    def q_header(self, vi, en):
        self._p(f"Các câu hỏi ôn tập đề mục – {vi} – {en}", "13. Importance")

    def qa(self, question_vi, answer_vi):
        self._p(question_vi, "01. GachTruDauDongViet")
        self._p(answer_vi, "01. GachTruDauDong")

    # -- end-of-chapter review questions ------------------------------------
    def review_section(self, questions):
        self._p("Các câu hỏi ôn tập – Questions for Review", "Heading 2")
        for q in questions:
            self._p(q, "01. GachTruDauDongViet")

    def _drop_orphan_images(self):
        """Remove image parts inherited from the template that the new body no
        longer references, so the output file stays small."""
        embed = qn("r:embed")
        link = qn("r:link")
        used = set()
        for blip in self.doc.element.body.iter(qn("a:blip")):
            for attr in (embed, link):
                rid = blip.get(attr)
                if rid:
                    used.add(rid)
        part = self.doc.part
        for rid, rel in list(part.rels.items()):
            if "image" in rel.reltype and rid not in used:
                part.drop_rel(rid)

    def _fix_chapter_footer(self):
        """The cloned template footer carries the static text 'Chương 15'. Replace
        it with a STYLEREF field so the footer shows the chapter of each page."""
        import re
        for section in self.doc.sections:
            for footer in (section.footer, section.first_page_footer,
                           section.even_page_footer):
                for p in footer.paragraphs:
                    for run in p.runs:
                        if re.match(r"^\s*Chương\s+\d+", run.text or ""):
                            r = run._r
                            for t in r.findall(qn("w:t")):
                                r.remove(t)
                            self._append_styleref(r)

    def _append_styleref(self, r):
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' STYLEREF "Heading 1" \\n \\* MERGEFORMAT '
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        cached = OxmlElement("w:t")
        cached.text = "Chương 1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for el in (begin, instr, sep, cached, end):
            r.append(el)

    def save(self, out_path):
        self._fix_chapter_footer()
        self._enable_update_fields()
        self._drop_orphan_images()
        self.doc.save(out_path)
        print("saved", out_path)
