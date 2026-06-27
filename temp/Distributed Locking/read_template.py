#!/usr/bin/env python3
"""Script to read and analyze the template DOCX file."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os

doc = Document("Fundamentals of Project Management Viet Anh.docx")

# Extract all paragraphs with style info
print("=" * 80)
print("DOCUMENT STRUCTURE AND CONTENT")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    style_name = para.style.name if para.style else "None"
    alignment = str(para.alignment) if para.alignment else "None"
    
    # Check for highlights and formatting
    formatting_info = []
    for run in para.runs:
        run_info = {
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_size': str(run.font.size) if run.font.size else None,
            'font_name': run.font.name,
            'font_color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
            'highlight_color': str(run.font.highlight_color) if run.font.highlight_color else None,
        }
        formatting_info.append(run_info)
    
    print(f"\n--- Paragraph {i} [Style: {style_name}] [Align: {alignment}] ---")
    print(f"Text: {text[:200]}")
    if formatting_info:
        for ri, info in enumerate(formatting_info):
            if any([info['bold'], info['italic'], info['underline'], info['highlight_color'], info['font_color']]):
                print(f"  Run {ri}: bold={info['bold']}, italic={info['italic']}, underline={info['underline']}, "
                      f"highlight={info['highlight_color']}, color={info['font_color']}, "
                      f"font={info['font_name']}, size={info['font_size']}")
                print(f"    Text: '{info['text'][:100]}'")

# Check for tables
print("\n" + "=" * 80)
print("TABLES IN DOCUMENT")
print("=" * 80)

for t_idx, table in enumerate(doc.tables):
    print(f"\n--- Table {t_idx} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for r_idx, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip()[:50])
        print(f"  Row {r_idx}: {' | '.join(row_data)}")

# Check for images
print("\n" + "=" * 80)
print("IMAGES IN DOCUMENT")
print("=" * 80)

from docx.opc.constants import RELATIONSHIP_TYPE as RT

image_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        image_count += 1
        print(f"  Image {image_count}: {rel.target_ref}")

print(f"\nTotal images found: {image_count}")

# Check sections
print("\n" + "=" * 80)
print("SECTIONS")
print("=" * 80)

for i, section in enumerate(doc.sections):
    print(f"Section {i}:")
    print(f"  Page width: {section.page_width}")
    print(f"  Page height: {section.page_height}")
    print(f"  Left margin: {section.left_margin}")
    print(f"  Right margin: {section.right_margin}")
    print(f"  Top margin: {section.top_margin}")
    print(f"  Bottom margin: {section.bottom_margin}")
